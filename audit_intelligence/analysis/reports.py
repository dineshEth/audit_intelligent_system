from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.shared import Inches

from ..utils.dates import utcnow_iso


class ReportBuilder:
    def __init__(self, settings) -> None:
        self.settings = settings

    def build_audit_report(
        self,
        source_name: str,
        summary_text: str,
        metrics: Dict[str, Any],
        references: List[Dict[str, Any]],
        chart_paths: Iterable[str] | None = None,
    ) -> str:
        stem = Path(source_name).stem
        path = self.settings.reports_dir / f"{stem}_audit_report_{utcnow_iso()}.docx"

        document = Document()
        document.add_heading("Audit Intelligence Report", level=0)
        document.add_paragraph(f"Source file: {source_name}")
        document.add_paragraph(f"Author: {self.settings.report_author}")

        document.add_heading("Executive Summary", level=1)
        document.add_paragraph(summary_text or "No summary available.")

        document.add_heading("Key Metrics", level=1)
        if metrics:
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Metric"
            table.rows[0].cells[1].text = "Value"
            for key, value in metrics.items():
                if isinstance(value, (dict, list)):
                    continue
                row = table.add_row().cells
                row[0].text = str(key)
                row[1].text = str(value)
        else:
            document.add_paragraph("No metrics were generated.")

        if metrics.get("anomalies"):
            document.add_heading("Flagged Transactions", level=1)
            for item in metrics["anomalies"]:
                document.add_paragraph(
                    f"{item.get('date')} | {item.get('description')} | debit={item.get('debit')} credit={item.get('credit')}",
                    style="List Bullet",
                )

        if references:
            document.add_heading("References", level=1)
            for idx, ref in enumerate(references, start=1):
                score = ref.get("score", "")
                snippet = ref.get("text") or ref.get("snippet") or ""
                document.add_paragraph(f"[{idx}] score={score} {snippet}", style="List Number")

        chart_paths = list(chart_paths or [])
        if chart_paths:
            document.add_heading("Charts", level=1)
            for chart in chart_paths:
                try:
                    document.add_picture(chart, width=Inches(6.2))
                except Exception:
                    document.add_paragraph(f"Chart available at: {chart}")

        document.save(path)
        return str(path)
