from __future__ import annotations

from pathlib import Path
from typing import Dict, List

import pandas as pd
from docx import Document
from docx.shared import Inches

from ..utils.dates import utcnow_iso
from ..utils.files import dump_dataframe_csv, dump_json


class LabelingExporter:
    def __init__(self, settings) -> None:
        self.settings = settings

    def export(self, source_name: str, labeled_df: pd.DataFrame, records: List[dict]) -> Dict[str, str]:
        stem = Path(source_name).stem
        timestamp = utcnow_iso()

        json_path = self.settings.labeled_data_dir / f"{stem}_labeled_{timestamp}.json"
        csv_path = self.settings.labeled_data_dir / f"{stem}_labeled_{timestamp}.csv"
        readable_csv_path = self.settings.labeled_docs_dir / f"{stem}_labeled_{timestamp}.csv"
        docx_path = self.settings.labeled_docs_dir / f"{stem}_labeled_{timestamp}.docx"

        dump_json(json_path, records)
        dump_dataframe_csv(csv_path, labeled_df)
        dump_dataframe_csv(readable_csv_path, labeled_df)
        self._build_docx(source_name, labeled_df, docx_path)

        return {
            "json": str(json_path),
            "csv": str(csv_path),
            "human_csv": str(readable_csv_path),
            "docx": str(docx_path),
        }

    def _build_docx(self, source_name: str, labeled_df: pd.DataFrame, path: Path) -> Path:
        document = Document()
        document.add_heading("Labeled Bank Statement", level=0)
        document.add_paragraph(f"Source file: {source_name}")
        document.add_paragraph(f"Rows labeled: {len(labeled_df)}")
        if "CATEGORY" in labeled_df.columns:
            category_counts = labeled_df["CATEGORY"].value_counts().to_dict()
            document.add_paragraph("Category counts: " + ", ".join(f"{k}={v}" for k, v in category_counts.items()))

        table = document.add_table(rows=1, cols=7)
        headers = ["DATE", "DESCRIPTION", "DEBIT", "CREDIT", "BALANCE", "CATEGORY", "CONFIDENCE"]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header

        preview = labeled_df[headers].head(25).fillna("")
        for _, row in preview.iterrows():
            cells = table.add_row().cells
            for idx, header in enumerate(headers):
                cells[idx].text = str(row[header])

        document.save(path)
        return path
